from docx import Document
from docx.shared import Pt, Mm
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE
from docx.oxml.shared import OxmlElement, qn
from docxtpl import DocxTemplate

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

# Функция выравнивания содержимого таблицы по вертикали
def set_cell_vertical_alignment(cell, align="center"): 
    try:   
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')  
        tcValign.set(qn('w:val'), align)  
        tcPr.append(tcValign)
        return True 
    except:
        #traceback.print_exc()             
        return False


def Word(fname,nam,pe,lv,t,project,nam_vl):
    doc = DocxTemplate("Word_template/Template.docx")
    context = { 'var_project' : project, "var_vl" : nam_vl}
    doc.render(context)
    doc.save("Word_template/generate.docx")

    # Создаём объект документа
    document = Document("Word_template/generate.docx")

    trig = False
    TextDown =[]
    for docpara in document.paragraphs:
        if trig:
            TextDown.append(docpara.text)
            delete_paragraph(docpara)
        elif docpara.text == "[table]":
            delete_paragraph(docpara)
            trig = True

    #print(TextDown)

    # Наследуем стиль и изменяем его
    style = document.styles['Normal'] # Берём стиль Нормальный

    f0 = style.font # Переменная для изменения параметров стиля
    #f0.name = 'Arial' # Шрифт
    f0.size = Pt(14) # Размер шрифта

    """ pf = style.paragraph_format
    pf.line_spacing = Pt(0) # Междустрочный интервал
    pf.space_after = Pt(0) # Интервал после абзаца

    # Задаём свойство полей
    sections = document.sections
    s = sections[0]
    s.left_margin = Mm(30)
    s.right_margin = Mm(15)
    s.top_margin = Mm(20)
    s.bottom_margin = Mm(20)
    s.page_height = Mm(297)
    s.page_width = Mm(210) """

    
    # Создаём шапку таблицы
    ptable=document.add_table(rows=1, cols=4, style='Table Grid') # Таблица с размерами и стилем
    
    p_t=ptable.style.paragraph_format
    #p_t.keep_with_next = True # Если невмещается таблица на текущую страницу, переносит её на ноыую
    
    row_1=ptable.rows[0] # Обращаемся к строке заголовка

    row_1.cells[0].text = 'Название опоры' # Присваиваем текст
    r = row_1.cells[0].paragraphs[0].runs[0] # Делаем его жирным
    row_1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_vertical_alignment(row_1.cells[0], align="center") # Выравнивание содержимого по вертикали
    #r.font.bold =True

    row_1.cells[1].text = 'Глубина заложения горизонтального заземлителя, м'
    row_1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_vertical_alignment(row_1.cells[1], align="center") # Выравнивание содержимого по вертикали
    r = row_1.cells[1].paragraphs[0].runs[0]
    #r.font.bold =True

    row_1.cells[2].text = 'Длинна вертикального заземлителя, м'
    row_1.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_vertical_alignment(row_1.cells[2], align="center") # Выравнивание содержимого по вертикали
    r = row_1.cells[2].paragraphs[0].runs[0]
    #r.font.bold =True

    row_1.cells[3].text = 'Удельное эквивалентное сопротивление грунта, Ом м'
    row_1.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_vertical_alignment(row_1.cells[3], align="center") # Выравнивание содержимого по вертикали
    r = row_1.cells[2].paragraphs[0].runs[0]
    #r.font.bold =True

    for i in range(len(nam)):         
        row_cells = ptable.add_row().cells # Добавляем в таблицу новую строку

        row_cells[0].text = str(nam[i])
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(row_cells[0], align="center") # Выравнивание содержимого по вертикали

        row_cells[1].text = str(t[i])
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(row_cells[1], align="center") # Выравнивание содержимого по вертикали

        row_cells[2].text = str(lv[i])
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(row_cells[2], align="center") # Выравнивание содержимого по вертикали

        row_cells[3].text = str(pe[i])
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(row_cells[3], align="center") # Выравнивание содержимого по вертикали

    
    document.add_paragraph("")
    for i in TextDown:
        document.add_paragraph(i)

    document.save(fname)

